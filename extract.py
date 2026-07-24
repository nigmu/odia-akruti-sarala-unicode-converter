"""Extract plain text from PDF and DOCX files for the conversion pipeline.

Why this file exists:
    pipeline.py needs a single entry point to load document text before Akruti
    Sarala to Unicode conversion. Input files may be PDF or DOCX, so format
    detection and extraction live here rather than in the pipeline orchestrator.

What it does:
    Detects file type from extension or file signature, then extracts text from
    PDF (PyMuPDF) or DOCX (python-docx) and returns it as a plain string.
"""

from __future__ import annotations

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx"})


class UnsupportedFormatError(ValueError):
    """Raised when a file is not a supported PDF or DOCX document."""


def detect_format(path: Path) -> str:
    """Return '.pdf' or '.docx' using the file extension, then magic bytes."""
    suffix = path.suffix.lower()
    if suffix in SUPPORTED_EXTENSIONS:
        return suffix

    with path.open("rb") as handle:
        header = handle.read(8)

    if header.startswith(b"%PDF"):
        return ".pdf"

    # DOCX files are ZIP archives (PK\x03\x04).
    if len(header) >= 2 and header[:2] == b"PK":
        return ".docx"

    raise UnsupportedFormatError(
        f"Unsupported file format for {path.name!r}. Expected .pdf or .docx."
    )


def extract_text_from_docx(path: str | Path) -> str:
    """Extract paragraph and table text from a DOCX file."""
    doc = Document(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = "".join(run.text for run in paragraph.runs)
        parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = "".join(run.text for run in paragraph.runs)
                    parts.append(text)

    return "\n".join(parts)


def extract_text_from_pdf(path: str | Path) -> str:
    """Extract text from each page of a PDF file."""
    with fitz.open(str(path)) as document:
        pages = [page.get_text("text") for page in document]
    return "\n".join(pages)


def extract_text(path: str | Path) -> str:
    """Extract text from a PDF or DOCX file, auto-detecting the format."""
    file_path = Path(path)
    if not file_path.is_file():
        raise FileNotFoundError(f"Input file not found: {file_path}")

    file_format = detect_format(file_path)
    if file_format == ".pdf":
        return extract_text_from_pdf(file_path)
    if file_format == ".docx":
        return extract_text_from_docx(file_path)

    raise UnsupportedFormatError(
        f"Unsupported file format for {file_path.name!r}. Expected .pdf or .docx."
    )
