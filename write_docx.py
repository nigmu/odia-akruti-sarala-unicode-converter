"""Write converted Unicode Odia text to a DOCX file.

Why this file exists:
    After Akruti Sarala text is converted to Unicode, the pipeline must save
    the result as a Word document that displays Odia correctly.

What it does:
    Creates a new DOCX with one paragraph per input line and applies a Unicode
    Odia font so the output renders properly in Word and other editors.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DEFAULT_ODIA_FONT = "Noto Sans Oriya"


def _apply_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def write_docx(
    text: str,
    output_path: str | Path,
    font_name: str = DEFAULT_ODIA_FONT,
) -> Path:
    """Write Unicode text to a DOCX file and return the output path."""
    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    lines = text.splitlines()
    if not lines and text:
        lines = [text]

    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        _apply_font(run, font_name)

    document.save(out_path)
    return out_path
